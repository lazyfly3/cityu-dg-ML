#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
读取所有assignment solutions的docx文件内容
"""
from docx import Document
import os
import sys

# 设置输出编码为UTF-8
if sys.platform == 'win32':
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

def read_docx(file_path):
    """读取docx文件内容"""
    try:
        doc = Document(file_path)
        content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text)
        
        # 也读取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    content.append(" | ".join(row_text))
        
        return "\n".join(content)
    except Exception as e:
        return f"读取文件时出错: {str(e)}"

# 读取所有solutions文件
files = [
    "Home_Assignment_1_solutions.docx",
    "Home_Assignment_2_solutions.docx",
    "Home_Assignment_3_solutions.docx",
    "Home_Assignment_4_solutions.docx"
]

# 将内容保存到文本文件
for filename in files:
    if os.path.exists(filename):
        output_filename = filename.replace('.docx', '.txt')
        content = read_docx(filename)
        with open(output_filename, 'w', encoding='utf-8') as f:
            f.write(f"{'='*80}\n")
            f.write(f"文件: {filename}\n")
            f.write(f"{'='*80}\n\n")
            f.write(content)
            f.write("\n\n")
        print(f"已读取并保存: {filename} -> {output_filename}")
    else:
        print(f"文件不存在: {filename}")

print("\n所有文件已读取完成！")

